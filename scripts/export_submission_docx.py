from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


AUTHOR_NAME = "Dr Siddalingaiah H S"
AUTHOR_AFFILIATION = "Professor, Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur"
AUTHOR_EMAIL = "hssling@yahoo.com"
AUTHOR_PHONE = "+91 8941087719"
AUTHOR_ORCID = "0000-0002-4771-8285"
TARGET_JOURNAL = "Tuberculosis (Elsevier)"


def strip_markdown(text: str) -> str:
    cleaned = text.replace("**", "").replace("`", "")
    cleaned = cleaned.replace("*", "")
    return cleaned.strip()


def add_table_from_markdown(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [strip_markdown(cell) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return

    header = rows[0]
    body = [row for row in rows[2:] if len(row) == len(header)]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for index, value in enumerate(header):
        table.rows[0].cells[index].text = value
    for row in body:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_markdown_image(document: Document, line: str, root: Path) -> bool:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return False
    caption, image_path = match.groups()
    image_file = (root / image_path).resolve()
    if not image_file.exists():
        return False
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_file), width=Inches(6.2))
    if caption:
        cap = document.add_paragraph(strip_markdown(caption))
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def markdown_to_docx(source: Path, target: Path, title: str) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table_from_markdown(document, table_buffer)
            table_buffer = []

    for raw_line in source.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_table()
            document.add_paragraph("")
            continue

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue

        flush_table()

        if add_markdown_image(document, stripped, source.parent):
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(strip_markdown(stripped[2:]))
            run.bold = True
            run.font.size = Pt(14)
            continue

        if stripped.startswith("## "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(strip_markdown(stripped[3:]))
            run.bold = True
            run.font.size = Pt(13)
            continue

        if stripped.startswith("### "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(strip_markdown(stripped[4:]))
            run.bold = True
            run.font.size = Pt(12)
            continue

        if stripped.startswith("- "):
            document.add_paragraph(strip_markdown(stripped[2:]), style="List Bullet")
            continue

        if stripped == "---":
            document.add_page_break()
            continue

        document.add_paragraph(strip_markdown(stripped))

    flush_table()
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def create_cover_letter(target: Path, manuscript_title: str) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    header = document.add_paragraph()
    header_run = header.add_run(AUTHOR_NAME)
    header_run.bold = True
    document.add_paragraph(AUTHOR_AFFILIATION)
    document.add_paragraph(f"Email: {AUTHOR_EMAIL}")
    document.add_paragraph(f"Phone: {AUTHOR_PHONE}")
    document.add_paragraph(f"ORCID: {AUTHOR_ORCID}")
    document.add_paragraph("")
    document.add_paragraph("The Editor-in-Chief")
    document.add_paragraph("Tuberculosis")
    document.add_paragraph("Elsevier")
    document.add_paragraph("")
    document.add_paragraph("Dear Editor,")
    document.add_paragraph(
        f"I am pleased to submit the manuscript titled \"{manuscript_title}\" for consideration as an Original Article in Tuberculosis."
    )
    document.add_paragraph(
        "India carries the world's largest tuberculosis burden, with approximately 119,000 MDR-TB cases annually and a treatment success rate of only 59% — figures that underscore the urgent need for new therapeutic options. InhA (enoyl-acyl carrier protein reductase, Rv1484) is the enzymatic target left intact by the dominant isoniazid-resistance mutation (katG S315T), making it the most rational focus for direct-inhibitor discovery in MDR-TB."
    )
    document.add_paragraph(
        "This study assembles 277 real, experimentally measured InhA IC50 values from ChEMBL, trains validated QSAR classifiers achieving ROC-AUC 0.961, and identifies CHEMBL3125270 (IC50 4 nM; pyrazole-benzofuran-pyrrolidine; MW 423 Da, cLogP 1.71, Lipinski-compliant) as the highest-priority computationally prioritized direct InhA inhibitor. No synthetic data, random imputation, or fabricated computational results were used at any stage. All analysis is fully reproducible from open-access scripts and publicly available databases."
    )
    document.add_paragraph(
        "This work falls squarely within the scope of Tuberculosis: it addresses the molecular basis of MDR-TB drug resistance, employs a reproducible computational workflow grounded in published bioactivity data, and proposes specific compounds for prospective experimental follow-up in the TB drug discovery pipeline. The transparent reporting of computational limitations — including the absence of molecular docking and molecular dynamics — is itself a contribution to the field's scientific integrity."
    )
    document.add_paragraph(
        "I confirm that this manuscript has not been published previously, is not under consideration for publication elsewhere, and all co-authors have approved the manuscript and agree with its submission. There are no competing interests to declare and no external funding was received."
    )
    document.add_paragraph("Thank you for considering this manuscript.")
    document.add_paragraph("")
    document.add_paragraph("Yours sincerely,")
    document.add_paragraph("")
    document.add_paragraph("Sincerely,")
    document.add_paragraph(AUTHOR_NAME)

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> None:
    parser = argparse.ArgumentParser(description="Export submission-ready DOCX assets.")
    parser.add_argument("--output-dir", default="outputs/submission", help="Directory for generated DOCX files.")
    args = parser.parse_args()

    root = Path(__file__).resolve().parent.parent
    output_dir = root / args.output_dir
    manuscript_md = root / "outputs" / "manuscript" / "manuscript_v8_genuine.md"
    supplementary_md = root / "outputs" / "manuscript" / "supplementary_materials.md"

    manuscript_title = strip_markdown(manuscript_md.read_text(encoding="utf-8").splitlines()[0].lstrip("# ").strip())

    # Use _v9 suffix to avoid overwriting any file open in Word
    markdown_to_docx(manuscript_md, output_dir / "submission_manuscript_v10.docx", manuscript_title)
    markdown_to_docx(supplementary_md, output_dir / "submission_supplementary_v10.docx", "Supplementary Materials")
    create_cover_letter(output_dir / "submission_cover_letter_v10.docx", manuscript_title)

    print(f"Generated DOCX assets in: {output_dir}")


if __name__ == "__main__":
    main()
